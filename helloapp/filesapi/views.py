# imageapi/views.py
from django.http import HttpResponse
from django.template.defaultfilters import slugify
import os
import PyPDF2
from docx import Document
from reportlab.pdfgen import canvas
import textwrap
import json
from django.http import HttpRequest
from django.shortcuts import render
from django.http import JsonResponse
from django.views.generic import TemplateView
from rest_framework.decorators import api_view
import mimetypes
from django.contrib.staticfiles.templatetags.staticfiles import static

# Create your views here.
class ConvertTXT(TemplateView):
    def get(self, request, **kwargs):
        form = int(request.GET['formate'])
        fileNames = request.GET['fileName']
        #txt to pdf
        if form == 500:
            c = canvas.Canvas(os.path.abspath("resources/static/downloads/"+fileNames+".pdf"))
            alltest = ""
            filename = os.path.abspath("resources/static/upload/file_215951.txt")
            for line in open(filename):
                alltest = alltest + line
            y = 227
            for lines in textwrap.wrap(alltest, 100):
                c.drawString(20, y, lines)
                y += 30

            c.showPage()
            c.save()
        # txt to docx
        elif form == 600:
            alltext = ""
            fileName = os.path.abspath("resources/static/upload/file_215951.txt")
            for line in open(fileName):
                alltext = alltext + line
            document = Document()
            document.add_paragraph(alltext)
            document.save(os.path.abspath('resources/static/downloads/'+fileNames+'.docx'))
        return JsonResponse({"status":HttpResponse().status_code,'response':"COMPLETE"})

class ConvertPDF(TemplateView):
    def get(self, request, **kwargs):
        pdf_file = open(os.path.abspath("resources/static/upload/file_21407.pdf"), "rb")
        read_pdf = PyPDF2.PdfFileReader(pdf_file)
        number_of_pages = read_pdf.getNumPages()
        page = ""
        form = int(request.GET['formate'])
        for v in range(number_of_pages):
            page = page+read_pdf.getPage(v).extractText()

        allContent = page
        fileName = request.GET['fileName']
        removeSpecialChars = allContent.translate({ord(c): " " for c in "!@#$%^&*()[]{};:,./<>?\|`~-=_+˙ˇ˙ˇˇ˜˜  ˙  ˝ ˙˜˘ ˆ  ˛ ˚ ˆˆˆ'  ˆ "})
        if form == 100:
            text_file = open(os.path.abspath("resources/static/downloads/"+fileName+".txt"), "w+")
            text_file.write("Conversion: %s" % removeSpecialChars)
            text_file.close()
        elif form == 200:
            document = Document()
            document.add_paragraph(removeSpecialChars)
            document.save(os.path.abspath("resources/static/downloads/"+fileName+".docx"))
        return JsonResponse({"status":HttpResponse().status_code,'response':"COMPLETE"})

class ConvertDOCX(TemplateView):
    def get(self, request, **kwargs):
        form = int(request.GET['formate'])
        fileName = request.GET['fileName']
        if form == 300:
            allContentOfDoc = ''
            document = Document(os.path.abspath('resources/static/upload/file_23525.docx'))
            for para in document.paragraphs:
                removeSpecialCharText = allContentOfDoc + para.text

            removeSpecialCharText = removeSpecialCharText.translate({ord(c): " " for c in "!@#$%^&*()[]{};:,./<>?\|`~-=_+˙ˇ˙ˇˇ˜˜  ˙  ˝ ˙˜˘ ˆ  ˛ ˚ ˆˆˆ'  ˆ "})
            text_file = open(os.path.abspath("resources/static/downloads/"+fileName+".txt"), "w+")
            text_file.write("Conversion: %s" % removeSpecialCharText)
            text_file.close()
        elif form == 400:
            from reportlab.lib.units import inch
            c = canvas.Canvas(os.path.abspath("resources/static/downloads/"+fileName+".pdf"))
            allContentOfDoc2 = ''
            addNewLine = ''
            document = Document(os.path.abspath('resources/static/upload/file_23525.docx'))
            for para in document.paragraphs:
                removeSpecialCharText2 = allContentOfDoc2 + para.text

            y = 227
            for line in textwrap.wrap(removeSpecialCharText2, 100):
                c.drawString(20, y, line)
                y += 30

            c.save()
        return JsonResponse({"status":HttpResponse().status_code,'response':"COMPLETE"})
